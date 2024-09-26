import os
import requests
import io

from aiogram.types import ReplyKeyboardRemove
from API_dwnld_file import get_download_url
import reply

from aiogram.types import BufferedInputFile, URLInputFile
from aiogram import  Bot, types, F, Router
from user_private import user_private_router
from calculating_functions import main_menu_back
from calculating_functions import buttom_back
from docx import Document
from aiogram.filters import CommandStart
from aiogram.filters import  StateFilter
import reply
from aiogram.fsm.state import State,StatesGroup
from aiogram.fsm.context import FSMContext



bot = Bot(token=os.getenv('TOKEN'))

class edit_docx(StatesGroup):
    awaiting_company_name = State()
    awaiting_plus1 = State()
    awaiting_plus2 = State()
    awaiting_plus3 = State()
    awaiting_choose1 = State()
    awaiting_choose2 = State()
    awaiting_choose3 = State()
    awaiting_output = State()
    awaiting_number = State()
    awaiting_email = State()
    awaiting_address = State()
    texts = {
    'edit_docx:awaiting_company_name': 'Отлично! Теперь введите название вашей компании.',
    'edit_docx:awaiting_choose1': "Преимущества компании сохранены ✅\n Хотите добавить еще преимущество?",
    'edit_docx:awaiting_choose2': "Преимущества компании сохранены ✅\n Хотите добавить еще преимущество?",
    'edit_docx:awaiting_plus1': "Название компании сохранено✅\nТеперь одним словом или одной фразой опишите преимущества компании:",
    'edit_docx:awaiting_plus2': "Одним словом или одной фразой опишите преимущества компании:",
    'edit_docx:awaiting_plus3': "Одним словом или одной фразой опишите преимущества компании:",
    'edit_docx:awaiting_number': "Преимущества компании сохранены ✅\nТеперь перейдем к заполнению контактов о компании.\nУкажите корпоративный телефон компании:",
    'edit_docx:awaiting_email': "Корпоративный номер сохранен ✅\nУкажите корпоративный email:",
    'edit_docx:awaiting_address': "Корпоративный email сохранен ✅\nУкажите корпоративный адрес:",
    'edit_docx:awaiting_output': 'Корпоративный адрес сохранен ✅',
}
    
def download_file(url):                                                             # Функция загрузки файла через url
    response = requests.get(get_download_url('disk:/листовка.docx'))
    return response.content

async def send_file(chat_id, file_content, filename):   # Функция выгрузки файла через url
    document = URLInputFile(file_content, filename=filename)
    # Отправка документа
    await bot.send_document(chat_id, document=document)

def replace_text_in_paragraph(paragraph, replacements):
    """Replace text in a single paragraph."""
    for old_string, new_string in replacements.items():
        if old_string in paragraph.text:
            paragraph.text = paragraph.text.replace(old_string, new_string)

def replace_text_in_run(run, replacements):
    """Replace text in a single run."""
    for old_string, new_string in replacements.items():
        if old_string in run.text:
            run.text = run.text.replace(old_string, new_string)

def replace_text_in_shapes(doc, replacements):
    """Replace text in all shapes in the document."""
    for shape in doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
        for paragraph in shape.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            for run in paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                for text in run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    for old_string, new_string in replacements.items():
                        if old_string in text.text:
                            text.text = text.text.replace(old_string, new_string)

def replace_text_in_tables(doc, replacements):
    """Replace text in all tables in the document."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

def replace_text_in_textboxes(doc, replacements):
    """Replace text in all textboxes in the document."""
    for shape in doc.inline_shapes:
        if not shape._inline.graphic.graphicData.uri.endswith("chart"):
            try:
                if shape.text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
            except AttributeError:
                continue

def edit_docx_file(content, replacements):
    doc = Document(io.BytesIO(content))
    
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Обработка текстовых полей (shapes)
    replace_text_in_shapes(doc, replacements)

    # Обработка таблиц
    replace_text_in_tables(doc, replacements)

    # Обработка текстовых полей (textboxes)
    replace_text_in_textboxes(doc, replacements)
    

    # Сохраняем изменения в документе
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()



async def send_edited_file(chat_id, file_content, filename="flyer.docx"):   # Функция для отправки редактированного файла пользователю
    # Создание объекта BufferedInputFile из данных файла в памяти
    document = BufferedInputFile(file_content, filename=filename)
    # Отправка документа
    await bot.send_document(chat_id, document=document)

@user_private_router.message(StateFilter('*'), F.text == '🔙 Вернуться в главное меню')
async def main_menu_back(message: types.Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state is None:
        return
    await message.answer('Выберите интересующее действие:', reply_markup=reply.start_keyboard)
    await state.clear()

    
@user_private_router.message(StateFilter(None), F.text == '✏️ Приступить к заполнению')
async def input_awaiting_company_name(message: types.Message, state: FSMContext):
    await message.answer("Отлично! Теперь введите название вашей компании.", reply_markup=reply.leave_keyboard)
    await state.set_state(edit_docx.awaiting_company_name)

@user_private_router.message(edit_docx.awaiting_company_name, F.text)
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await state.update_data(awaiting_company_name=message.text)
    await message.answer("Название компании сохранено✅\nТеперь одним словом или одной фразой опишите преимущества компании:", reply_markup=reply.back_kb)
    await state.set_state(edit_docx.awaiting_plus1)

@user_private_router.message(edit_docx.awaiting_plus1, F.text)
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await state.update_data(awaiting_plus1=message.text)
    await message.answer("Преимущества компании сохранены ✅\nХотите добавить еще преимущество?", reply_markup=reply.accept_keyboard)
    await state.set_state(edit_docx.awaiting_choose1)

@user_private_router.message(edit_docx.awaiting_choose1, F.text == '✅ Да')
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await message.answer("Одним словом или одной фразой опишите преимущества компании:", reply_markup=reply.back_kb)
    await state.set_state(edit_docx.awaiting_plus2)

@user_private_router.message(edit_docx.awaiting_choose1, F.text == '❌ Нет')
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await message.answer("Преимущества компании сохранены ✅\nТеперь перейдем к заполнению контактов о компании.\nУкажите корпоративный телефон компании:", reply_markup=ReplyKeyboardRemove())
    await state.set_state(edit_docx.awaiting_number)

@user_private_router.message(edit_docx.awaiting_plus2, F.text)
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await state.update_data(awaiting_plus2=message.text)
    await message.answer("Преимущества компании сохранены ✅\n Хотите добавить еще преимущество?", reply_markup=reply.accept_keyboard)
    await state.set_state(edit_docx.awaiting_choose2)

@user_private_router.message(edit_docx.awaiting_choose2, F.text == '✅ Да')
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await message.answer("Одним словом или одной фразой опишите преимущества компании:", reply_markup=reply.back_kb)
    await state.set_state(edit_docx.awaiting_plus3)

@user_private_router.message(edit_docx.awaiting_choose2, F.text == '❌ Нет')
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await message.answer("Преимущества компании сохранены ✅\nТеперь перейдем к заполнению контактов о компании.\nУкажите корпоративный телефон компании:", reply_markup=ReplyKeyboardRemove())
    await state.set_state(edit_docx.awaiting_number)

@user_private_router.message(edit_docx.awaiting_plus3, F.text)
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await state.update_data(awaiting_plus3=message.text)
    await message.answer("Преимущества компании сохранены ✅\nТеперь перейдем к заполнению контактов о компании.\nУкажите корпоративный телефон компании:", reply_markup=ReplyKeyboardRemove())
    await state.set_state(edit_docx.awaiting_number)

@user_private_router.message(edit_docx.awaiting_number, F.text)
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    if message.text.isdigit():
        await state.update_data(awaiting_number=message.text)
        await message.answer("Корпоративный номер сохранен ✅\nУкажите корпоративный email:", reply_markup=reply.back_kb)
        await state.set_state(edit_docx.awaiting_email)
    else:
        await message.answer("На этом этапе я понимаю только числа 😰\nУкажите корпоративный телефон компании:")

@user_private_router.message(edit_docx.awaiting_email, F.text)
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await state.update_data(awaiting_email=message.text)
    await message.answer("Корпоративный email сохранен ✅\nУкажите корпоративный адрес:", reply_markup=reply.back_kb)
    await state.set_state(edit_docx.awaiting_address)

@user_private_router.message(edit_docx.awaiting_address, F.text)
async def input_awaiting_competitors(message: types.Message, state: FSMContext):
    await state.update_data(awaiting_address=message.text)
    await message.answer("Корпоративный адрес сохранен ✅", reply_markup=reply.upload_file_keyboard)
    await state.set_state(edit_docx.awaiting_output)

@user_private_router.message(edit_docx.awaiting_output, F.text)
async def output(message: types.Message, state: FSMContext):
    chat_id = message.chat.id
    data = await state.get_data()
    file_content = download_file(get_download_url('disk:/листовка.docx'))
    replacements = {
        "НАЗВАНИЕ КОМПАНИИ" : data.get('awaiting_company_name', ''),
        "ПРЕИМУЩЕСТВО 1" : data.get('awaiting_plus1', ''),
        "ПРЕИМУЩЕСТВО 2" : data.get('awaiting_plus2', ''),
        "ПРЕИМУЩЕСТВО 3" : data.get('awaiting_plus3', ''),
        "ТЕЛЕФОН" : data.get('awaiting_number', ''),
        "ИМЭИЛ" : data.get('awaiting_email', ''),
        "ADR" : data.get('awaiting_address', ''),
                   }
    edited_content = edit_docx_file(file_content, replacements)
    await send_edited_file(chat_id, edited_content)
    await message.answer("Ваш флаер готов 😉", reply_markup=reply.start_keyboard)                                                 
    await state.clear()

